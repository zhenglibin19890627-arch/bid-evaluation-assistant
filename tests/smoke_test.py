# -*- coding: utf-8 -*-
"""合成数据冒烟测试：验证流水线关键路径（初筛/增量/多 Sheet 回写/状态双轨）。

用法：在项目根目录执行 `python tests/smoke_test.py`。
"""
import os
import sys
import json
import tempfile
import subprocess
from datetime import datetime, timedelta

PROJECT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT)

import pandas as pd
from scripts import run_pipeline, excel_tool

failures = []


def check(name, cond, detail=""):
    print(("PASS  " if cond else "FAIL  ") + name + ("  | " + detail if detail else ""))
    if not cond:
        failures.append(name)


base = tempfile.mkdtemp(prefix="bid_smoke_")
list_dir = os.path.join(base, "list")
bid_dir = os.path.join(base, "bid")
prof = os.path.join(base, "profile")
os.makedirs(list_dir), os.makedirs(bid_dir), os.makedirs(prof)

today = datetime.now()
dtstr = lambda d: d.strftime("%Y-%m-%d %H:%M")

rows = [
    [dtstr(today), "采购项目公告", "莲都区", "丽水市实验小学", "丽水市实验小学办公电脑采购项目",
     "100000", "", "", "浙江XX代理有限公司", "https://example.com/u1"],
    [dtstr(today), "采购结果公告", "莲都区", "丽水市实验小学", "丽水市实验小学课桌椅采购项目",
     "50000", "48000", "杭州某供应商有限公司", "浙江XX代理有限公司", "https://example.com/u2"],
    [dtstr(today), "电子卖场公告", "莲都区", "丽水市实验小学", "丽水市实验小学打印机采购",
     "20000", "", "", "", "https://example.com/u3"],
    [dtstr(today - timedelta(days=10)), "采购项目公告", "莲都区", "丽水市实验小学",
     "丽水市实验小学空调采购项目", "80000", "", "", "", "https://example.com/u4"],
]
cols = ["公告时间", "公告类型", "所在区县", "采购人", "采购项目名称",
        "预算金额(元)", "中标金额(元)", "中标单位", "代理机构", "地址链接URL"]
pd.DataFrame(rows, columns=cols).to_excel(os.path.join(list_dir, "丽水清单.xlsx"), index=False)

excel_tool.save_selected_types({"办公设备与家具": []}, path=os.path.join(list_dir, "项目类型选择.json"))
excel_tool.save_time_range(7, path=os.path.join(list_dir, "时间范围选择.json"))

md = """# 公司介绍
## 公司基本信息
- 公司名称：浙江云帆科技有限公司
- 注册资本：1000 万元
## 公司业务
- 主营业务：办公设备销售与售后
- 服务领域：政府、教育
## 资质证书
- ISO9001：持有
## 历史业绩
- 近三年：8 个同类项目，累计 1200 万元
"""
with open(os.path.join(prof, "公司介绍.md"), "w", encoding="utf-8") as f:
    f.write(md)

from docx import Document
doc = Document()
doc.add_heading("丽水市实验小学办公电脑采购项目 招标文件", 0)
doc.add_paragraph("采购人：丽水市实验小学")
doc.add_paragraph("预算金额：100000 元")
doc.add_heading("四、评分标准", level=1)
doc.add_paragraph("企业资质（ISO9001 认证，得 15 分，无则 0 分）")
doc.add_paragraph("业绩（近三年同类业绩≥3 个得满分 25 分，每少 1 个扣 5 分）")
doc.add_paragraph("人员（项目经理 1 名提供社保，得 20 分，否则 0 分）")
doc.save(os.path.join(bid_dir, "丽水市实验小学办公电脑采购项目招标文件.docx"))

# ---- 1) 时间不符文案随天数 ----
raw = pd.read_excel(os.path.join(list_dir, "丽水清单.xlsx"), dtype=str)
fdf = excel_tool.filter_screening(raw, company_text=md, days=7,
                                  project_types=excel_tool.load_project_types(os.path.join(list_dir, "项目类型选择.json")))
reason10 = fdf[fdf["地址链接URL"] == "https://example.com/u4"]["初筛原因"].iloc[0]
check("时间不符文案随窗口（7 天）", reason10 == "公告超过 7 天", reason10)

# ---- 2) 流水线全量运行 ----
r1 = run_pipeline.run_pipeline(list_dir, bid_dir, prof)
s1 = r1["summary"]
check("全量初筛：4 行中 1 行通过（电子卖场/类型/项目类型/时间过滤）", s1["passed_count"] == 1, str(s1["screening"]))
check("标书匹配成功", s1["matched_count"] == 1, r1["items"][0].get("bid_match_way", ""))
check("评分章节提取成功", not r1["items"][0]["scoring_missing"])
check("初筛结果 Excel 已生成", bool(r1["screening_excel"]) and os.path.exists(r1["screening_excel"]))
check("公司介绍无占位词告警", not any("示例" in e for e in r1["errors"]), "errors 数: %d" % len(r1["errors"]))

# ---- 3) 增量模式三轮 ----
r2 = run_pipeline.run_pipeline(list_dir, bid_dir, prof, incremental=True)
check("增量第 1 轮：全部为新增", r2["summary"].get("new_count") == 4, str(r2["summary"].get("new_count")))
r3 = run_pipeline.run_pipeline(list_dir, bid_dir, prof, incremental=True)
check("增量第 2 轮：全部跳过（new=0, skipped=4）",
      r3["summary"].get("new_count") == 0 and r3["summary"].get("skipped_screened_count") == 4,
      "new=%s skipped=%s" % (r3["summary"].get("new_count"), r3["summary"].get("skipped_screened_count")))

# ---- 4) 多 Sheet 主文件回写工作流 ----
list2 = os.path.join(base, "list2")
os.makedirs(list2)
import openpyxl
wb = openpyxl.Workbook()
sheet_rows = {
    "丽水": [rows[0], rows[3]],
    "杭州": [[rows[0][0], rows[0][1], "西湖区", rows[0][3], "杭州市某中学办公电脑采购项目",
              rows[0][5], "", "", "", "https://example.com/u5"],
             [dtstr(today - timedelta(days=10)), "采购项目公告", "西湖区", "杭州市某中学",
              "杭州市某中学空调采购项目", "60000", "", "", "", "https://example.com/u6"]],
}
for sn, srows in sheet_rows.items():
    ws = wb.create_sheet(sn)
    ws.append(cols + ["初筛状态", "研判状态", "初筛结果", "初筛原因", "相关性(命中词)"])
    for r in srows:
        ws.append(r + ["", "", "", "", ""])
wb.remove(wb["Sheet"])
main_path = os.path.join(list2, "政府采购公告.xlsx")
wb.save(main_path)
excel_tool.save_selected_types({"办公设备与家具": []}, path=os.path.join(list2, "项目类型选择.json"))
excel_tool.save_time_range(7, path=os.path.join(list2, "时间范围选择.json"))
r4 = run_pipeline.run_pipeline(list2, bid_dir, prof)
wb2 = openpyxl.load_workbook(main_path)
ws2 = wb2["丽水"]
col_of = {c.value: i for i, c in enumerate(ws2[1], 1)}
screened_vals = [ws2.cell(2, col_of["初筛结果"]).value, ws2.cell(3, col_of["初筛结果"]).value]
wb2.close()
check("多 Sheet 主文件：初筛结果回写主文件对应 Sheet、不新建结果表",
      r4["screening_excel"] == main_path and not any("初筛结果" in f for f in os.listdir(list2))
      and screened_vals == ["通过", "时间不符"],
      str(screened_vals))

# ---- 4b) status_tool 同 URL 跨 Sheet 状态回写 ----
list3 = os.path.join(base, "list3")
os.makedirs(list3)
wb = openpyxl.Workbook()
for sn in ("甲", "乙"):
    ws = wb.create_sheet(sn)
    ws.append(["公告时间", "公告类型", "采购项目名称", "地址链接URL", "初筛状态", "研判状态"])
    ws.append([dtstr(today), "采购项目公告", "项目X", "https://example.com/dup", "", ""])
wb.remove(wb["Sheet"])
dup_path = os.path.join(list3, "政府采购公告.xlsx")
wb.save(dup_path)
from scripts import status_tool
status_tool.write_state(list3, [{"url": "https://example.com/dup", "初筛状态": "已初筛", "研判状态": "已研判"}])
wb3 = openpyxl.load_workbook(dup_path)
vals = [(wb3["甲"].cell(2, 5).value, wb3["甲"].cell(2, 6).value),
        (wb3["乙"].cell(2, 5).value, wb3["乙"].cell(2, 6).value)]
wb3.close()
check("status_tool 同 URL 跨 Sheet 均回写", vals == [("已初筛", "已研判"), ("已初筛", "已研判")], str(vals))

# ---- 5) analyze_main 主文件缺失时落盘 ----
ws_dir = os.path.join(base, "ws_empty")
os.makedirs(ws_dir)
subprocess.run([sys.executable, os.path.join(PROJECT, "scripts", "analyze_main.py"),
                "--city", "丽水市", "--workspace", ws_dir], capture_output=True, cwd=PROJECT)
ar_path = os.path.join(ws_dir, "analysis_result.json")
if os.path.exists(ar_path):
    with open(ar_path, encoding="utf-8") as f:
        ar = json.load(f)
    check("analyze_main 主文件缺失时落盘 3 天", len(ar.get("incomplete_dates", [])) == 3,
          str(ar.get("incomplete_dates")))
else:
    check("analyze_main 主文件缺失时落盘 3 天", False, "json 未生成")

print("\n临时工作区：%s" % base)
print("RESULT: %s" % ("ALL PASS" if not failures else "FAILED: " + "、".join(failures)))
sys.exit(1 if failures else 0)
